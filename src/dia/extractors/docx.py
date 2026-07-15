"""DOCX text extraction using python-docx."""

import logging
from io import BytesIO

import docx

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extracts text and tables from DOCX documents."""

    def extract(self, content: bytes) -> str:
        """Extract text and tables from DOCX bytes.

        Paragraphs are extracted first, followed by tables formatted
        as markdown pipe tables.
        """
        file_stream = BytesIO(content)
        content_parts: list[str] = []

        doc = docx.Document(file_stream)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            content_parts.append("--- Document Text ---")
            content_parts.extend(paragraphs)

        if doc.tables:
            content_parts.append("\n--- Document Tables ---")
            for table in doc.tables:
                content_parts.append(self._format_table(table))
                content_parts.append("\n")

        return "\n".join(content_parts).strip()

    def _format_table(self, table) -> str:
        """Format a docx table as a markdown pipe table."""
        rows = []
        for row in table.rows:
            try:
                rows.append(self._format_row(row))
            except Exception:
                logger.warning("Skipping malformed table row", exc_info=True)
                continue
        return "\n".join(rows)

    def _format_row(self, row) -> str:
        """Format a single row as a markdown pipe row."""
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        return "| " + " | ".join(cells) + " |"
